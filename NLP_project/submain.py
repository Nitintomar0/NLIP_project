import tempfile
from PIL import Image
!pip install PyMuPDF
!pip install python-docx
!pip install gradio
!pip install gradio -q
import nltk
nltk.download('punkt_tab', quiet=True) # Download the punkt_tab resource
nltk.download('punkt', quiet=True)  # Correct
import tempfile
from PIL import Image


import os
import re
import nltk
import spacy
import fitz
import docx
import numpy as np
import seaborn as sns
import matplotlib.pyplot as plt
import gradio as gr

from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from wordcloud import WordCloud
from sentence_transformers import SentenceTransformer

nlp = spacy.load('en_core_web_sm')
# In the cell where NLTK is first imported or used, add this line
def clean_text(text):
    text = re.sub(r'[^\w\s]', '', text.lower())
    return text

def preprocess(text):
    doc = nlp(clean_text(text))
    return ' '.join([token.lemma_ for token in doc if not token.is_stop])

model = SentenceTransformer('all-MiniLM-L6-v2')  # Lightweight and fast

def cosine_sim(text1, text2):
    vectorizer = TfidfVectorizer()
    vecs = vectorizer.fit_transform([text1, text2])
    return cosine_similarity(vecs[0:1], vecs[1:2])[0][0]

def bert_sim(text1, text2):
    embeddings = model.encode([text1, text2])
    return cosine_similarity([embeddings[0]], [embeddings[1]])[0][0]

def show_wordcloud(text):
    wordcloud = WordCloud(width=800, height=400, background_color='white').generate(text)
    fig, ax = plt.subplots()
    ax.imshow(wordcloud, interpolation='bilinear')
    ax.axis('off')
    return fig

def similarity_bar(sim_dict):
    fig, ax = plt.subplots()
    sns.barplot(x=list(sim_dict.keys()), y=list(sim_dict.values()), ax=ax)
    ax.set_title("Text Similarity Scores")
    ax.set_ylabel("Score")
    ax.set_ylim(0, 1)
    return fig

# Remove the following lines as the functions are defined in the notebook
# from similarity import cosine_sim, bert_sim
# from utils import preprocess
# from visualizer import similarity_bar, show_wordcloud

def analyze(text1, text2):
    p1, p2 = preprocess(text1), preprocess(text2)
    cosine = cosine_sim(p1, p2)
    bert = bert_sim(text1, text2)

    # Generate visualizations
    wordcloud_path = dual_wordclouds(p1, p2)
    dist_plot_path = similarity_distribution_plot(text1, text2)
    heatmap_path = sentence_similarity_heatmap(text1, text2)

    return f"Cosine Similarity: {cosine:.2f}\nBERT Similarity: {bert:.2f}", wordcloud_path, dist_plot_path, heatmap_path



def wrapper(file1, file2):
    text1 = extract_text(file1)
    text2 = extract_text(file2)
    return analyze(text1, text2)

iface = gr.Interface(
    fn=wrapper,
    inputs=[gr.File(label="Upload File 1"), gr.File(label="Upload File 2")],
    outputs=[
        gr.Textbox(label="Similarity Scores"),
        gr.Image(type="filepath", label="Word Clouds"),
        gr.Image(type="filepath", label="Similarity Distribution"),
        gr.Image(type="filepath", label="Sentence Heatmap")
    ],
    title="Document Similarity Analyzer with Visuals"
)

def extract_text(file):

    ext = os.path.splitext(file.name)[1]
    if ext == ".pdf":
        import fitz
        doc = fitz.open(file.name)
        return " ".join([page.get_text() for page in doc])
    elif ext == ".docx":
        import docx
        doc = docx.Document(file.name)
        return "\n".join([para.text for para in doc.paragraphs])
    elif ext == ".txt":
        return file.read().decode("utf-8")
    else:
        return ""
def chunk_and_embed(text, model, chunk_size=512):

    from transformers import BertTokenizer
    tokenizer = BertTokenizer.from_pretrained('bert-base-uncased')

    sentences = nltk.sent_tokenize(text)
    chunks = []
    current_chunk = ""

    for sent in sentences:
        if len(tokenizer.encode(current_chunk + sent)) <= chunk_size:
            current_chunk += " " + sent
        else:
            chunks.append(current_chunk.strip())
            current_chunk = sent
    if current_chunk:
        chunks.append(current_chunk.strip())

    embeddings = model.encode(chunks)
    return np.mean(embeddings, axis=0)  # Average over all chunk embeddings

def split_into_paragraphs(text, min_length=30):
    """Split text into non-empty paragraphs"""
    paragraphs = [p.strip() for p in text.split('\n') if len(p.strip()) > min_length]
    return paragraphs

def paragraph_level_similarity(text1, text2):
    paras1 = split_into_paragraphs(text1)
    paras2 = split_into_paragraphs(text2)

    embeddings1 = model.encode(paras1)
    embeddings2 = model.encode(paras2)

    sim_matrix = cosine_similarity(embeddings1, embeddings2)
    return paras1, paras2, sim_matrix

def show_similarity_heatmap(sim_matrix, title='Similarity Heatmap'):
    plt.figure(figsize=(10, 8))
    sns.heatmap(sim_matrix, annot=True, cmap='YlGnBu', fmt=".2f")
    plt.title(title)
    plt.xlabel("Document 2 Paragraphs")
    plt.ylabel("Document 1 Paragraphs")
    plt.tight_layout()
    plt.show()

from IPython.display import display, HTML

def render_highlighted_sentences(highlights):
    html = "<h3>Similar Sentences</h3><ul>"
    for s1, s2, score in highlights:
        html += f"<li><span style='color:green;'>{s1}</span><br><span style='color:blue;'>{s2}</span><br><b>Score:</b> {score:.2f}</li><br>"
    html += "</ul>"
    display(HTML(html))
    ext = os.path.splitext(file.name)[1]
    if ext == ".pdf":
        import fitz
        doc = fitz.open(file.name)
        return " ".join([page.get_text() for page in doc])
    elif ext == ".docx":
        import docx
        doc = docx.Document(file.name)
        return "\n".join([para.text for para in doc.paragraphs])
    elif ext == ".txt":
        return file.read().decode("utf-8")
    else:
        return ""
def chunk_and_embed(text, model, chunk_size=512):

    from transformers import BertTokenizer
    tokenizer = BertTokenizer.from_pretrained('bert-base-uncased')

    sentences = nltk.sent_tokenize(text)
    chunks = []
    current_chunk = ""

    for sent in sentences:
        if len(tokenizer.encode(current_chunk + sent)) <= chunk_size:
            current_chunk += " " + sent
        else:
            chunks.append(current_chunk.strip())
            current_chunk = sent
    if current_chunk:
        chunks.append(current_chunk.strip())

    embeddings = model.encode(chunks)
    return np.mean(embeddings, axis=0)  # Average over all chunk embeddings


iface.launch()

def similarity_distribution_plot(text1, text2):
    sents1 = nltk.sent_tokenize(text1)
    sents2 = nltk.sent_tokenize(text2)
    emb1 = model.encode(sents1)
    emb2 = model.encode(sents2)

    scores = [cosine_similarity([e1], [e2])[0][0] for e1 in emb1 for e2 in emb2]

    fig, ax = plt.subplots()
    sns.histplot(scores, bins=20, kde=True, ax=ax, color='skyblue')
    ax.set_title("Similarity Score Distribution")
    ax.set_xlabel("Score")
    ax.set_ylabel("Frequency")

    temp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    fig.savefig(temp.name)
    plt.close(fig)
    return temp.name
def dual_wordclouds(text1, text2):
    wordcloud1 = WordCloud(width=400, height=300, background_color='white').generate(text1)
    wordcloud2 = WordCloud(width=400, height=300, background_color='white').generate(text2)

    fig, axs = plt.subplots(1, 2, figsize=(12, 6))
    axs[0].imshow(wordcloud1, interpolation='bilinear')
    axs[0].axis('off')
    axs[0].set_title('Document 1')

    axs[1].imshow(wordcloud2, interpolation='bilinear')
    axs[1].axis('off')
    axs[1].set_title('Document 2')

    plt.tight_layout()

    temp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    fig.savefig(temp.name)
    plt.close(fig)
    return temp.name
def sentence_similarity_heatmap(text1, text2):
    sents1 = nltk.sent_tokenize(text1)
    sents2 = nltk.sent_tokenize(text2)
    emb1 = model.encode(sents1)
    emb2 = model.encode(sents2)
    sim_matrix = cosine_similarity(emb1, emb2)

    fig, ax = plt.subplots(figsize=(10, 8))
    sns.heatmap(sim_matrix, cmap="coolwarm", xticklabels=sents2, yticklabels=sents1)
    ax.set_title("Sentence Similarity Heatmap")
    plt.xticks(rotation=90)
    plt.yticks(rotation=0)
    plt.tight_layout()

    temp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    fig.savefig(temp.name)
    plt.close(fig)
    return temp.name

def highlight_similar_sentences(text1, text2, threshold=0.75):
    sents1 = nltk.sent_tokenize(text1)
    sents2 = nltk.sent_tokenize(text2)

    emb1 = model.encode(sents1)
    emb2 = model.encode(sents2)

    highlights = []
    for i, e1 in enumerate(emb1):
        for j, e2 in enumerate(emb2):
            sim = cosine_similarity([e1], [e2])[0][0]
            if sim > threshold:
                highlights.append((sents1[i], sents2[j], sim))
    return highlights

def render_highlighted_sentences(highlights):
    html = "<h3>Similar Sentences</h3><ul>"
    for s1, s2, score in highlights:
        html += f"<li><span style='color:green;'>{s1}</span><br><span style='color:blue;'>{s2}</span><br><b>Score:</b> {score:.2f}</li><br>"
    html += "</ul>"
    display(HTML(html))